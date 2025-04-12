from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile

def nota_letra(puntos):
    if puntos >= 9: return 'A', puntos
    elif puntos >= 7: return 'B', puntos
    elif puntos >= 5: return 'C', puntos
    else: return 'D', puntos

def calcular_nota(tics):
    puntos = (tics / 6) * 10
    return nota_letra(round(puntos, 1))

def generar_informe(datos, conceptos, observaciones_generales):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Verdana'
    font.size = Pt(10)

    p = doc.add_paragraph("INFORME PERSONAL DE LA GUARDIA")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].bold = True

    for k in ["informante", "fecha", "alumno", "puesto"]:
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table.style = "Table Grid"
        cells = table.rows[0].cells
        cells[0].text = k.upper()
        cells[1].text = datos[k]
        for c in cells:
            c.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()

    tabla = doc.add_table(rows=1 + len(conceptos), cols=6)
    tabla.style = 'Table Grid'
    tabla.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tabla.autofit = False

    headers = ["CONCEPTOS A CALIFICAR", "A", "B", "C", "D", "OBSERVACIONES"]
    for i, text in enumerate(headers):
        cell = tabla.rows[0].cells[i]
        cell.text = text
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    notas = []
    for i, concepto in enumerate(conceptos):
        nota_letra_, nota_num = calcular_nota(concepto["ticks"])
        notas.append(nota_num)
        row = tabla.rows[i + 1].cells
        row[0].text = concepto["categoria"]
        for j, l in enumerate(["A", "B", "C", "D"]):
            row[j + 1].text = "X" if l == nota_letra_ else ""
            row[j + 1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row[5].text = f"({nota_num}) {concepto['observacion']}"
        row[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        row[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()
    doc.add_paragraph("NOTAS: A (De 9 a 10), B (de 7 a 8.9), C (de 5 a 6.9), D (menos de 5)").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()

    media = round(sum(notas) / len(notas), 1)
    media_letra = nota_letra(media)[0]
    fila_media = doc.add_table(rows=1, cols=2)
    fila_media.style = 'Table Grid'
    fila_media.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fila_media.rows[0].cells[0].text = "Nota media:"
    fila_media.rows[0].cells[1].text = f"{media} ({media_letra})"
    for c in fila_media.rows[0].cells:
        c.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()
    doc.add_paragraph("OBSERVACIONES GENERAL / JUSTIFICACIÓN:")
    doc.add_paragraph(observaciones_generales).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    output = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(output)
    return output
